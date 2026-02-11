#!/usr/bin/env python3
"""
Generate comprehensive DOCX reports from expanded markdown files.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

BASE_DIR = "/Users/yerassyl/dev/ML_masters_degree"

def add_image_if_exists(doc, img_path, width=5.5, caption=None):
    """Add image to document if it exists."""
    full_path = os.path.join(BASE_DIR, img_path) if not os.path.isabs(img_path) else img_path
    if os.path.exists(full_path):
        doc.add_picture(full_path, width=Inches(width))
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.italic = True
            run.font.size = Pt(10)
    else:
        doc.add_paragraph(f"[Изображение не найдено: {os.path.basename(img_path)}]")

def parse_md_to_docx(md_path, docx_path):
    """Parse markdown file and generate DOCX."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    doc = Document()
    
    # Main title if first line is # header
    if lines and lines[0].startswith('# '):
        title = lines[0][2:].strip()
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lines = lines[1:]  # Skip first line
    
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # End of code block
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.5)
                in_code_block = False
                code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Tables
        if line.startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # Process table
            if len(table_lines) > 2:
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                # Skip separator line
                rows = []
                for tl in table_lines[2:]:
                    row = [cell.strip() for cell in tl.split('|')[1:-1]]
                    if row and any(row):
                        rows.append(row)
                
                if headers and rows:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    hdr_cells = table.rows[0].cells
                    for idx, h in enumerate(headers):
                        hdr_cells[idx].text = h
                        hdr_cells[idx].paragraphs[0].runs[0].bold = True
                    for row in rows:
                        row_cells = table.add_row().cells
                        for idx, val in enumerate(row):
                            if idx < len(row_cells):
                                row_cells[idx].text = val
            in_table = False
            table_lines = []
            # Don't increment i, process current line
            continue
        
        # Headings
        if line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        
        # Images
        elif line.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                caption, img_path = match.groups()
                add_image_if_exists(doc, img_path, 5.5, caption)
        
        # Horizontal rule
        elif line.startswith('---') or line.startswith('***'):
            doc.add_paragraph()
        
        # Skip certain formatting lines
        elif line.startswith('*Рис') or line.startswith('**Ключевые'):
            p = doc.add_paragraph(line)
            p.runs[0].font.italic = True
        
        # Regular paragraph
        elif line.strip():
            # Skip pure bold lines (often section markers)
            if not (line.startswith('**') and line.endswith('**')):
                # Handle inline formatting
                doc.add_paragraph(line)
        
        i += 1
    
    doc.save(docx_path)
    print(f"Saved: {docx_path}")
    return docx_path

if __name__ == '__main__':
    # Generate main report
    parse_md_to_docx(
        os.path.join(BASE_DIR, 'Анализ и разработка методов сегментации сцен и 3D реконструкции по данным глубины_expanded.md'),
        os.path.join(BASE_DIR, 'Анализ и разработка методов сегментации сцен и 3D реконструкции по данным глубины.docx')
    )
    
    # Generate technical report
    parse_md_to_docx(
        os.path.join(BASE_DIR, 'ТЕХНИЧЕСКИЙ ОТЧЕТ_expanded.md'),
        os.path.join(BASE_DIR, 'ТЕХНИЧЕСКИЙ ОТЧЕТ.docx')
    )
    
    print("\nВсе отчёты успешно сгенерированы!")
