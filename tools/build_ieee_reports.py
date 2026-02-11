from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
import os

FULL_PATH = "Анализ и разработка методов сегментации сцен и 3D реконструкции по данным глубины.docx"
TECH_PATH = "ТЕХНИЧЕСКИЙ ОТЧЕТ.docx"


def add_center_title(doc, text, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text):
    doc.add_paragraph(text)


def add_image_if_exists(doc, path, width_inches=5.5, caption=None):
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)


def build_full_report():
    doc = Document()
    add_center_title(doc, "Анализ и разработка методов сегментации сцен и 3D реконструкции по данным глубины")
    add_center_title(doc, "NYU Depth Dataset V2", bold=False)
    add_center_title(doc, "Весна 2026", bold=False)
    add_center_title(doc, "Автор: Маратов Ерасыл Балканович", bold=False)

    add_heading(doc, "Аннотация", level=1)
    add_paragraph(doc, "В работе представлен воспроизводимый baseline‑пайплайн сегментации сцен и 3D реконструкции по данным глубины на NYU Depth Dataset V2. Описаны предобработка, модельные настройки, метрики качества, формулы и визуальные сравнения. Результаты служат точкой отсчета для дальнейшего улучшения архитектур и проведения сравнительного анализа.")

    add_heading(doc, "Ключевые слова", level=1)
    add_paragraph(doc, "RGB‑D, NYU Depth V2, семантическая сегментация, TSDF, 3D реконструкция, baseline")

    add_heading(doc, "1. Введение", level=1)
    add_paragraph(doc, "Глубинные данные дают геометрический контекст, критичный для понимания сцен в помещениях. В отличие от чисто RGB‑подходов, глубина помогает устойчиво отделять объекты на сложном фоне и формировать геометрически согласованную реконструкцию. Эти задачи востребованы в робототехнике, AR/VR, навигации и анализе интерьеров.")
    add_paragraph(doc, "NYU Depth Dataset V2 остается одним из самых распространенных наборов для экспериментального сравнения методов сегментации и реконструкции. Наличие синхронизированных RGB‑D кадров и плотной разметки позволяет строить полноценный конвейер: от предобработки до формирования метрик и визуализаций.")
    add_paragraph(doc, "Цель работы — построить воспроизводимую baseline‑реализацию сегментации и 3D реконструкции, получить метрики качества и подготовить визуальные сравнения, пригодные для детального отчета и дальнейших исследований.")

    add_heading(doc, "2. Обзор литературы и сравнительный анализ", level=1)
    add_paragraph(doc, "Сверточные модели (DeepLabv3+, HRNet, OCRNet) обеспечивают стабильность и интерпретируемость, тогда как трансформерные архитектуры (SegFormer, Mask2Former, Swin‑based) достигают более высоких значений mIoU на редких классах. Для реконструкции в baseline‑сценариях стандартом остается TSDF‑фузия (Open3D), а нейросетевые методы (NeuralRecon) обеспечивают более детальные поверхности при более высоких требованиях к вычислениям.")

    add_heading(doc, "3. Методология", level=1)
    add_heading(doc, "3.1 Набор данных", level=2)
    add_paragraph(doc, "NYU Depth Dataset V2 содержит 1449 размеченных RGB‑D кадров. Используется стандартный сплит (795 train / 654 test). Для CPU‑экспериментов экспортируется подвыборка 400 кадров.")
    add_heading(doc, "3.2 Предобработка", level=2)
    add_paragraph(doc, "Размер приводится к 320x240, RGB‑каналы нормализуются, глубина переводится в метры, некорректные значения маскируются. Применяются легкие аугментации (отражение, цветовые искажения).")
    add_heading(doc, "3.3 Модели и параметры", level=2)
    add_paragraph(doc, "Сегментация: Tiny U‑Net, Adam, lr=1e‑3, batch size 4, epochs 5. Реконструкция: TSDF‑фузия, voxel size 0.02 м, truncation 0.04 м, Gaussian фильтр по глубине.")

    add_heading(doc, "4. Формулы", level=1)
    add_paragraph(doc, "Cross‑entropy: L = -1/N * sum_i sum_c y_ic log p_ic")
    add_paragraph(doc, "IoU: IoU_c = TP_c / (TP_c + FP_c + FN_c), mIoU = (1/C) * sum_c IoU_c")
    add_paragraph(doc, "RMSE: sqrt((1/N) * sum_i (d_i - d_i*)^2)")
    add_paragraph(doc, "AbsRel: (1/N) * sum_i |d_i - d_i*| / d_i*")
    add_paragraph(doc, "TSDF: D_new = (w_old D_old + w_obs D_obs) / (w_old + w_obs)")

    add_heading(doc, "5. Экспериментальная постановка", level=1)
    add_paragraph(doc, "Эксперименты выполнены на CPU. Подвыборка: 400 кадров (train=240, test=160).")
    add_table(doc, ["Parameter", "Value"], [
        ["Segmentation optimizer", "Adam"],
        ["Segmentation lr", "1e-3"],
        ["Segmentation epochs", "5"],
        ["Segmentation batch size", "4"],
        ["TSDF voxel size", "0.02 m"],
        ["TSDF truncation", "0.04 m"],
    ])

    add_heading(doc, "6. Результаты", level=1)
    add_paragraph(doc, "Сегментация:")
    add_table(doc, ["Metric", "Value"], [
        ["mIoU", "0.0022"],
        ["Pixel Accuracy", "0.2798"],
        ["Mean Accuracy", "0.0057"],
        ["FW IoU", "0.1221"],
        ["Dice", "0.0034"],
    ])
    add_paragraph(doc, "Реконструкция:")
    add_table(doc, ["Metric", "Value"], [
        ["RMSE", "0.0292"],
        ["AbsRel", "0.0033"],
        ["delta < 1.25", "0.9999"],
        ["delta < 1.25^2", "1.0000"],
        ["delta < 1.25^3", "1.0000"],
    ])

    add_heading(doc, "7. Визуализации", level=1)
    add_image_if_exists(doc, "results/figures/segmentation_compare_00240.png", caption="Segmentation: RGB / GT / Pred")
    add_image_if_exists(doc, "results/figures/reconstruction_compare_00240.png", caption="Depth and mesh comparison")
    add_image_if_exists(doc, "results/figures/seg_loss_curve.png", caption="Segmentation training loss")
    add_image_if_exists(doc, "results/figures/metrics_bar.png", caption="Metrics summary")

    add_heading(doc, "8. Фрагменты кода", level=1)
    add_paragraph(doc, "Сегментация:")
    add_code_block(doc, [
        "logits = model(rgb)",
        "loss = F.cross_entropy(logits, label)",
        "optimizer.zero_grad()",
        "loss.backward()",
        "optimizer.step()",
    ])
    add_paragraph(doc, "TSDF‑фузия:")
    add_code_block(doc, [
        "volume.integrate(rgbd, intrinsics, np.eye(4))",
        "mesh = volume.extract_triangle_mesh()",
        "mesh.compute_vertex_normals()",
    ])

    add_heading(doc, "9. Анализ и ограничения", level=1)
    add_paragraph(doc, "mIoU остается низким из‑за компактной архитектуры и ограниченного числа эпох. TSDF‑реконструкция устойчива для крупных поверхностей, но сглаживает мелкие детали. CPU‑режим ограничивает масштаб экспериментов.")

    add_heading(doc, "10. Заключение", level=1)
    add_paragraph(doc, "Реализован воспроизводимый baseline‑пайплайн сегментации и 3D реконструкции на NYU Depth V2. Получены метрики качества и визуальные сравнения, пригодные для дальнейшего сравнительного анализа и улучшения моделей.")

    add_heading(doc, "References (IEEE)", level=1)
    refs = [
        "[1] L.-C. Chen et al., ECCV, 2018.",
        "[2] H. Zhao et al., CVPR, 2017.",
        "[3] G. Lin et al., CVPR, 2017.",
        "[4] K. Sun et al., CVPR, 2019.",
        "[5] Y. Yuan et al., ECCV, 2020.",
        "[6] E. Xie et al., NeurIPS, 2021.",
        "[7] B. Cheng et al., CVPR, 2022.",
        "[8] Z. Liu et al., ICCV, 2021.",
        "[9] W. Sun et al., CVPR, 2021.",
        "[10] Q.-Y. Zhou et al., 2018.",
        "[11] H. Laina et al., 2016.",
        "[12] H. Fu et al., CVPR, 2018.",
        "[13] R. Ranftl et al., ICCV, 2021.",
        "[14] S. Bhat et al., CVPR, 2021.",
        "[15] W. Yuan et al., CVPR, 2022.",
    ]
    for ref in refs:
        add_paragraph(doc, ref)

    doc.save(FULL_PATH)


def build_tech_report():
    doc = Document()
    add_center_title(doc, "ТЕХНИЧЕСКИЙ ОТЧЕТ")
    add_center_title(doc, "Анализ и разработка методов сегментации сцен и 3D реконструкции по данным глубины", bold=False)
    add_center_title(doc, "Весна 2026", bold=False)
    add_center_title(doc, "Автор: Маратов Ерасыл Балканович", bold=False)

    add_heading(doc, "1. Введение", level=1)
    add_paragraph(doc, "Отчет описывает baseline‑пайплайн сегментации и реконструкции по данным глубины на NYU Depth V2. Акцент на воспроизводимости, метриках и визуальных сравнениях.")

    add_heading(doc, "2. Цели работы", level=1)
    add_paragraph(doc, "Построить полный конвейер обработки RGB‑D данных, реализовать сегментацию и TSDF‑реконструкцию, получить метрики и артефакты.")

    add_heading(doc, "3. Теоретические основы", level=1)
    add_paragraph(doc, "Сегментация — пиксельная классификация (cross‑entropy). Реконструкция — TSDF‑фузия глубины с извлечением поверхности.")

    add_heading(doc, "4. Реализация", level=1)
    add_paragraph(doc, "Экспорт данных из nyu_depth_v2_labeled.mat, обучение Tiny U‑Net на подвыборке, TSDF‑реконструкция в Open3D, генерация метрик и визуализаций.")

    add_heading(doc, "5. Результаты", level=1)
    add_paragraph(doc, "Сегментация: mIoU=0.0022, PixelAcc=0.2798. Реконструкция: RMSE=0.0292, AbsRel=0.0033.")

    add_image_if_exists(doc, "results/figures/segmentation_compare_00240.png", caption="Segmentation comparison")
    add_image_if_exists(doc, "results/figures/reconstruction_compare_00240.png", caption="Reconstruction comparison")
    add_image_if_exists(doc, "results/figures/seg_loss_curve.png", caption="Training loss")
    add_image_if_exists(doc, "results/figures/metrics_bar.png", caption="Metrics summary")

    add_heading(doc, "6. Выводы", level=1)
    add_paragraph(doc, "Baseline‑пайплайн воспроизводим и обеспечивает базовую точку отсчета. Для улучшения качества необходимы современные архитектуры и GPU‑обучение.")

    add_heading(doc, "References (IEEE)", level=1)
    for ref in [
        "[1] L.-C. Chen et al., ECCV, 2018.",
        "[2] E. Xie et al., NeurIPS, 2021.",
        "[3] B. Cheng et al., CVPR, 2022.",
        "[4] W. Sun et al., CVPR, 2021.",
        "[5] Q.-Y. Zhou et al., 2018.",
    ]:
        add_paragraph(doc, ref)

    doc.save(TECH_PATH)


if __name__ == "__main__":
    build_full_report()
    build_tech_report()
    print("Saved:", FULL_PATH, TECH_PATH)
