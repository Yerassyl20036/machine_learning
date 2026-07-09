"""Generate iris recognition project report in docx format."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def make_doc():
    doc = Document()

    # page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(1.5)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)

    def h(text, level=1):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.bold = True
            run.font.size = Pt(14 if level == 1 else 13)
            run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def para(text, bold=False, center=False, italic=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def cp(text, bold=False, size=12, italic=False):
        """Centered paragraph shortcut."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        return p

    def add_table(headers, rows, col_widths=None):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, hdr in enumerate(headers):
            hdr_cells[i].text = hdr
            r = hdr_cells[i].paragraphs[0].runs[0]
            r.font.bold = True
            r.font.size = Pt(11)
            r.font.name = 'Times New Roman'
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = str(val)
                for r in cell.paragraphs[0].runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
        if col_widths:
            for trow in table.rows:
                for i, w in enumerate(col_widths):
                    if i < len(trow.cells):
                        trow.cells[i].width = Cm(w)
        return table

    # ── TITLE PAGE ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    cp("АСТАНА ХАЛЫҚАРАЛЫҚ УНИВЕРСИТЕТІ", bold=True, size=14)
    cp("Ақпараттық технологиялар факультеті", size=12)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    cp("ЖОБА ЕСЕБІ", bold=True, size=16)
    doc.add_paragraph()
    cp("Тақырыбы: Көздің мұқабасын тану әдістері мен алгоритмдері", bold=True, size=14)
    cp("(Методы и алгоритмы распознавания радужной оболочки глаза)", italic=True, size=13)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    cp("Орындаған: Маратов Ерасыл Балканович", size=12)
    cp("1 курс магистранты", size=12)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    cp("Астана, 2025", size=12)
    doc.add_page_break()

    # ── 1. ВВЕДЕНИЕ ─────────────────────────────────────────────────────────
    h("1. Кіріспе (Введение)")
    para(
        "Биометриялық аутентификация — заманауи ақпараттық қауіпсіздіктің маңызды бөлігі. "
        "Дактилоскопия, бет-жүз тану және дауыс тану сияқты модальдіктер арасында "
        "радужная оболочка (iris) ерекше сенімділікпен ерекшеленеді: жоғары энтропия, "
        "тұрақтылық және генетикалық айрықшалық оны биометриядағы «алтын стандарт» "
        "деп санауға мүмкіндік береді."
    )
    para(
        "Осы жобада классикалық алгоритмдер (Daugman IrisCode, LBP, HOG, ORB) мен "
        "терең оқыту моделі (IrisNet CNN) салыстырылады. Деректер жиыны ретінде "
        "университет оқушыларынан жиналған 16 адамның суреттері пайдаланылды."
    )

    # ── 2. ЦЕЛИ ─────────────────────────────────────────────────────────────
    h("2. Мақсаты мен міндеттері (Цели и задачи)")
    para("Мақсаты:", bold=True)
    para("Радужная оболочкасын тану үшін бірнеше алгоритмді іске асыру, салыстыру және ансамбльдеу.")
    para("Міндеттері:", bold=True)
    for t in [
        "1. Iris деректер жиынын жинау және алдын ала өңдеу (нормализация, сегментация).",
        "2. Daugman IrisCode, LBP, HOG, ORB алгоритмдерін іске асыру.",
        "3. Weighted Ensemble және Stacking Ensemble ансамбльдерін құру.",
        "4. IrisNet CNN моделін PyTorch негізінде оқыту.",
        "5. Алгоритмдерді EER, AUC, Accuracy метрикалары бойынша салыстыру.",
        "6. Нақты уақыттағы демонстрация скриптін жазу.",
    ]:
        para(t)

    # ── 3. СТРУКТУРА ────────────────────────────────────────────────────────
    h("3. Жоба құрылымы (Структура проекта)")
    para("Жобаның файлдық құрылымы:")
    add_table(
        ["Файл / Қалта", "Сипаттама"],
        [
            ["main.py",                    "Негізгі орындаушы скрипт (benchmark pipeline)"],
            ["iris_demo.py",               "Демонстрациялық CLI утилита"],
            ["requirements.txt",           "Python тәуелділіктері"],
            ["src/preprocess.py",          "Кескінді алдын ала өңдеу (CLAHE, Hough)"],
            ["src/dataset.py",             "Деректер жиынын жүктеу"],
            ["src/metrics.py",             "EER, AUC, ROC есептеу"],
            ["src/neural_model.py",        "IrisNet CNN архитектурасы"],
            ["src/ensemble.py",            "Weighted & Stacking Ensemble"],
            ["src/algorithms/daugman.py",  "Daugman IrisCode алгоритмі"],
            ["src/algorithms/lbp.py",      "LBP + косинустық ұқсастық"],
            ["src/algorithms/hog_algo.py", "HOG дескрипторы"],
            ["src/algorithms/orb_algo.py", "ORB keypoint matching"],
            ["results/irisnet.pt",         "Оқытылған CNN модель (PyTorch)"],
            ["results/comparison/algorithm_comparison.csv", "Алгоритм салыстыру нәтижелері"],
            ["results/cnn_history.json",   "CNN оқыту тарихы"],
        ],
        col_widths=[7, 10],
    )
    doc.add_paragraph()

    # ── 4. ДАННЫЕ ────────────────────────────────────────────────────────────
    h("4. Деректер сипаттамасы (Описание данных)")
    para(
        "Деректер жиыны — Астана халықаралық университетінің 16 студентінен алынған "
        "радужная оболочка суреттері. Жалпы: 16 субъект × 80 сурет = 1 280 кескін."
    )
    add_table(
        ["Параметр", "Мән"],
        [
            ["Субъект саны",         "16 адам"],
            ["Бір адамға суреттер",  "80 кескін"],
            ["Жалпы кескін",         "1 280"],
            ["Кескін форматы",       "JPG / PNG"],
            ["Кескін өлшемі",        "320×280 пиксель (орт.)"],
            ["Таңбалар",             "Субъект аттары (классификация)"],
        ],
        col_widths=[7, 10],
    )
    doc.add_paragraph()

    # ── 5. БИБЛИОТЕКИ ────────────────────────────────────────────────────────
    h("5. Пайдаланылған кітапханалар (Используемые библиотеки)")
    add_table(
        ["Кітапхана", "Қолдану мақсаты"],
        [
            ["numpy",               "Сандық есептеулер"],
            ["opencv-python",       "Кескін өңдеу, Hough трансформациясы, CLAHE"],
            ["scikit-image",        "LBP, HOG дескрипторлары"],
            ["scikit-learn",        "SVM, Stacking Classifier, метрикалар"],
            ["torch / torchvision", "IrisNet CNN оқыту"],
            ["matplotlib",          "ROC қисықтары, тарих графиктері"],
            ["pandas",              "Нәтижелерді CSV-ге жазу"],
            ["python-docx",         "Есепті .docx форматта генерациялау"],
        ],
        col_widths=[5, 12],
    )
    doc.add_paragraph()

    # ── 6. МЕТОДЫ ────────────────────────────────────────────────────────────
    h("6. Әдістер және алгоритмдер (Методы и алгоритмы)")

    h("6.1 Алдын ала өңдеу (Предобработка)", level=2)
    para(
        "Әр кескін үшін: (1) сұр реңге аудару; (2) Hough дөңгелек трансформациясы "
        "арқылы қарашықты (pupil) және иристі (iris) локализациялау; "
        "(3) CLAHE контрастты жақсарту; (4) нормализация — polar координаттар "
        "жүйесіне айналдыру (Daugman rubber-sheet)."
    )

    h("6.2 Daugman IrisCode", level=2)
    para(
        "Нормализацияланған иристен 2D Gabor фильтрі арқылы бинарлық код алынады. "
        "Ұқсастық өлшемі — Hamming арақашықтығы:"
    )
    p = doc.add_paragraph("HD(A, B) = (A XOR B) · mask / N_mask")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.italic = True
    p.runs[0].font.name = 'Times New Roman'

    h("6.3 LBP (Local Binary Pattern)", level=2)
    para(
        "Кескіннен LBP гистограммасы алынады (P=8, R=1, uniform режим — 64 бин). "
        "Ұқсастық өлшемі — косинустық ұқсастық:"
    )
    p = doc.add_paragraph("cos(A, B) = (A · B) / (||A|| · ||B||)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.italic = True
    p.runs[0].font.name = 'Times New Roman'
    para("Шекті мән: 0.980  (sim >= 0.980  =>  MATCH).")

    h("6.4 HOG (Histogram of Oriented Gradients)", level=2)
    para(
        "Пиксельдер (8, 8) жасушаларға бөлінеді, градиент бағыты 9 бинге санақталады. "
        "Ұқсастық өлшемі — косинустық ұқсастық."
    )

    h("6.5 ORB (Oriented FAST + Rotated BRIEF)", level=2)
    para(
        "FAST детекторы мен BRIEF дескрипторының бұрылысқа тұрақты нұсқасы. "
        "Сәйкестілік — Hamming арақашықтығы, BFMatcher арқылы."
    )

    h("6.6 Weighted & Stacking Ensemble", level=2)
    para(
        "Weighted Ensemble: классификаторлар баллдары ең жақсы AUC-ке пропорционал "
        "салмақтармен қосылады. Stacking Ensemble: SVM мета-классификатор, "
        "бастапқы алгоритмдер деңгей-0 болады."
    )

    h("6.7 IrisNet CNN", level=2)
    para(
        "PyTorch негізіндегі конволюциялық нейрондық желі: "
        "Conv2d(3,32) -> Conv2d(32,64) -> MaxPool -> Dropout(0.3) -> "
        "FC(256) -> FC(16). Оқыту: 30 эпоха, CrossEntropyLoss, Adam optimizer."
    )

    # ── 7. НӘТИЖЕЛЕР ─────────────────────────────────────────────────────────
    h("7. Нәтижелер (Результаты)")

    h("7.1 Алгоритм салыстыру кестесі", level=2)
    add_table(
        ["Алгоритм", "Метрика", "EER (аз = жақсы)", "AUC (көп = жақсы)", "Accuracy"],
        [
            ["Daugman IrisCode",    "Hamming", "0.3375", "0.7197", "94.14%"],
            ["LBP",                 "Cosine",  "0.2812", "0.7821", "70.31%"],
            ["HOG",                 "Cosine",  "0.3437", "0.7146", "92.58%"],
            ["ORB",                 "Hamming", "0.3021", "0.7683", "75.78%"],
            ["Ensemble (Weighted)", "–",       "0.3000", "0.7713", "–"],
            ["Ensemble (Stacking)", "–",       "0.2625", "0.8298", "–"],
            ["IrisNet CNN",         "Softmax", "–",      "–",      "95.70%*"],
        ],
        col_widths=[5, 3, 4, 4, 3],
    )
    doc.add_paragraph()
    para("* IrisNet CNN — 30 эпоха оқытудың нәтижесі (ең жақсы validation accuracy).")

    h("7.2 Демонстрация нәтижелері (iris_demo.py)", level=2)
    para("iris_demo.py CLI скриптін іске қосу мысалдары (LBP Cosine, threshold=0.980):")
    add_table(
        ["Команда", "Ұқсастық", "Нәтиже", "Ескерту"],
        [
            ["compare Arnur_1 vs Arnur_5",   "0.9941", "MATCH",    "Бір адам — дұрыс"],
            ["compare Arnur_1 vs Aibar_1",   "0.9666", "NO MATCH", "Басқа адам — дұрыс"],
            ["batch /tmp/iris_test/ (6 жұп)", "–",     "3M / 3NM", "Дұрыс бөлу"],
        ],
        col_widths=[6.5, 3, 3, 4.5],
    )
    doc.add_paragraph()

    # ── 8. СУРЕТТЕР ───────────────────────────────────────────────────────────
    h("8. Суреттер тізімі (Список рисунков)")
    for fig in [
        "Сурет 1 — ROC қисығы: барлық алгоритмдер (results/figures/roc_curves.png)",
        "Сурет 2 — IrisNet CNN оқыту тарихы: loss & accuracy (results/figures/cnn_history.png)",
        "Сурет 3 — Алгоритм салыстыру: EER & AUC бар диаграмма (results/figures/comparison.png)",
    ]:
        para(fig)

    # ── 9. ҚОРЫТЫНДЫ ─────────────────────────────────────────────────────────
    h("9. Қорытынды (Заключение)")
    para(
        "Жобада 4 классикалық алгоритм (Daugman, LBP, HOG, ORB), 2 ансамбль "
        "(Weighted, Stacking) және IrisNet CNN модел сынақтан өтті. "
        "Ең жақсы EER (0.2625) — Stacking Ensemble, ал ең жақсы AUC (0.8298) "
        "да Stacking Ensemble-ге тиесілі. "
        "CNN жоғары accuracy (95.70%) көрсетті, алайда аз деректер кезінде "
        "overfitting байқалады."
    )
    para(
        "LBP алгоритмі қарапайым, интерпретациялануға ыңғайлы және нақты уақыттағы "
        "тануда (threshold = 0.980) сенімді нәтиже береді. Болашақ жұмыста "
        "трансформерлік архитектуралар мен деректер аугментациясын қолдану ұсынылады."
    )

    # ── 10. ӘДЕБИЕТТЕР ───────────────────────────────────────────────────────
    h("10. Пайдаланылған әдебиеттер (Список литературы)")
    for ref in [
        "1. Daugman J. How iris recognition works // IEEE Transactions on Information "
           "Forensics and Security. — 2004. — Vol. 1. — P. 21–31.",
        "2. Ojala T., Pietikainen M., Maenpaa T. Multiresolution Gray-Scale and "
           "Rotation Invariant Texture Classification with Local Binary Patterns // "
           "IEEE TPAMI. — 2002. — Vol. 24, No. 7. — P. 971–987.",
        "3. Dalal N., Triggs B. Histograms of Oriented Gradients for Human Detection // "
           "Proc. CVPR. — 2005.",
        "4. Rublee E. et al. ORB: An efficient alternative to SIFT or SURF // "
           "Proc. ICCV. — 2011.",
        "5. He K., Zhang X., Ren S., Sun J. Deep Residual Learning for Image "
           "Recognition // Proc. CVPR. — 2016.",
        "6. Pedregosa F. et al. Scikit-learn: Machine Learning in Python // "
           "JMLR. — 2011. — Vol. 12. — P. 2825–2830.",
        "7. Paszke A. et al. PyTorch: An Imperative Style, High-Performance "
           "Deep Learning Library // NeurIPS. — 2019.",
    ]:
        para(ref)

    out = "/Users/yerassyl/dev/ML_masters_degree/Маратов_Е_радужная_оболочка.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    make_doc()
