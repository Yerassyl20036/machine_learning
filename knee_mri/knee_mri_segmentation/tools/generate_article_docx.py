#!/usr/bin/env python3
"""
Generate the article (статья) docx for:
  Маратов Ерасыл Балканович
  Методы автоматической сегментации костных и хрящевых тканей коленного сустава по данным МРТ

Based on the format of example articles (Сакатаган, Абдуап).
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = Path(__file__).parent
FIGURES_DIR = BASE_DIR / "results" / "eda_figures"
HW_DIR = BASE_DIR / "results" / "homework_template_style"


def set_run(run, font_name="Times New Roman", font_size=14, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
                  font_size=14, space_after=Pt(6), space_before=Pt(0),
                  first_line_indent=Cm(1.0), italic=False):
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = space_before
    if first_line_indent:
        para.paragraph_format.first_line_indent = first_line_indent
    run = para.add_run(text)
    set_run(run, font_size=font_size, bold=bold, italic=italic)
    return para


def add_heading_text(doc, text, font_size=14):
    """Add a bold centered heading-like paragraph."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.first_line_indent = Cm(1.0)
    run = para.add_run(text)
    set_run(run, font_size=font_size, bold=True)
    return para


def add_figure(doc, image_path, caption, width=Inches(5.5)):
    """Add a figure with centered caption."""
    if not os.path.exists(image_path):
        add_paragraph(doc, f"[Рисунок не найден: {image_path}]", italic=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run()
    run.add_picture(str(image_path), width=width)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.first_line_indent = None
    run = cap.add_run(caption)
    set_run(run, font_size=14)


def add_table(doc, headers, rows, title=None):
    """Add a formatted table (Times New Roman 14) with optional title."""
    if title:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(4)
        cap.paragraph_format.first_line_indent = None
        run = cap.add_run(title)
        set_run(run, font_size=14)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, font_size=14, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run(run, font_size=14)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer


def build_article():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.0

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # =====================================================================
    # МРНТИ
    # =====================================================================
    mrnti = doc.add_paragraph()
    mrnti.alignment = WD_ALIGN_PARAGRAPH.LEFT
    mrnti.paragraph_format.space_after = Pt(6)
    mrnti.paragraph_format.first_line_indent = None
    run = mrnti.add_run("МРНТИ 28.23.15")
    set_run(run, font_size=14, bold=True)

    # =====================================================================
    # TITLE
    # =====================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.first_line_indent = None
    run = title.add_run(
        "Методы автоматической сегментации костных и хрящевых тканей "
        "коленного сустава по данным МРТ"
    )
    set_run(run, font_size=14, bold=True)

    # Author
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(2)
    author.paragraph_format.first_line_indent = None
    run = author.add_run("Маратов Е.Б.")
    set_run(run, font_size=14, bold=True)

    # University
    uni = doc.add_paragraph()
    uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    uni.paragraph_format.space_after = Pt(2)
    uni.paragraph_format.first_line_indent = None
    run = uni.add_run("Астана халықаралық университеті, Астана, Қазақстан")
    set_run(run, font_size=14)

    # Author description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    desc.paragraph_format.space_after = Pt(6)
    desc.paragraph_format.first_line_indent = Cm(1.0)
    run = desc.add_run(
        "Маратов Е.Б. – магистрант 1 курса, Астана халықаралық университеті, Астана, Казахстан"
    )
    set_run(run, font_size=14, bold=True)

    # =====================================================================
    # ABSTRACT
    # =====================================================================
    abs_para = doc.add_paragraph()
    abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_para.paragraph_format.space_after = Pt(6)
    abs_para.paragraph_format.first_line_indent = Cm(1.0)
    r1 = abs_para.add_run("Аннотация. ")
    set_run(r1, bold=True)
    r2 = abs_para.add_run(
        "В работе рассматривается задача автоматической сегментации костных и хрящевых тканей "
        "коленного сустава по данным магнитно-резонансной томографии (МРТ). "
        "В качестве источника данных использован открытый набор 3D Knee MRI Cartilage "
        "Segmentation, основанный на материалах Osteoarthritis Initiative (OAI), содержащий "
        "МРТ-срезы с попиксельной разметкой пяти анатомических классов: фон, бедренный хрящ, "
        "большеберцовый хрящ, костная ткань и дефекты хряща. Проведён разведочный анализ "
        "данных с извлечением 20 интерпретируемых признаков (интенсивностные, спектральные, "
        "текстурные GLCM, морфологические и клинические). Реализован базовый метод сегментации "
        "на основе порогового метода Otsu и морфологических операций (Dice 0.20–0.35). "
        "Построены и сравнены пять моделей классификации степени остеоартрита по шкале "
        "Kellgren–Lawrence: Logistic Regression, Random Forest, SVM, Decision Tree и KNN. "
        "Лучшая модель (Logistic Regression) достигла точности 75.22% (F1 = 75.19%). "
        "Также построена модель линейной регрессии для прогнозирования прогрессии остеоартрита "
        "с R² = 0.85 и RMSE = ±0.37%. Ключевыми предикторами являются ширина суставной щели, "
        "выраженность остеофитов и доля хрящевой ткани. Результаты подтверждают возможность "
        "применения методов машинного обучения для анализа МРТ-данных коленного сустава."
    )
    set_run(r2)

    # Keywords
    kw_para = doc.add_paragraph()
    kw_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_para.paragraph_format.space_after = Pt(10)
    kw_para.paragraph_format.first_line_indent = Cm(1.0)
    r1 = kw_para.add_run("Ключевые слова: ")
    set_run(r1, bold=True)
    r2 = kw_para.add_run(
        "сегментация коленного сустава, МРТ, остеоартрит, машинное обучение, "
        "линейная регрессия, Random Forest, метод Otsu"
    )
    set_run(r2)

    # =====================================================================
    # ВВЕДЕНИЕ
    # =====================================================================
    add_heading_text(doc, "Введение")

    add_paragraph(doc,
        "Остеоартрит (ОА) коленного сустава является одним из наиболее распространённых "
        "дегенеративных заболеваний опорно-двигательного аппарата, затрагивающим более "
        "250 миллионов человек по всему миру [18]. Магнитно-резонансная томография (МРТ) "
        "считается золотым стандартом визуализации мягких тканей сустава, позволяя "
        "оценивать состояние хрящевой ткани и костных структур без ионизирующего "
        "излучения."
    )

    add_paragraph(doc,
        "Однако ручная сегментация МРТ-изображений трудоёмка: один МРТ-объём содержит "
        "40–160 срезов. Архитектуры U-Net [1], nnU-Net [2] и TransUNet [3] достигают "
        "Dice > 0.85, но остаются проблемы: дисбаланс классов, размытые границы тканей "
        "на поздних стадиях ОА и потребность в комплексном анализе."
    )

    p = add_paragraph(doc, "")
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.0)
    r1 = p.add_run("Цель исследования: ")
    set_run(r1, bold=True)
    r2 = p.add_run(
        "разработка и оценка методов автоматической сегментации и анализа "
        "МРТ-изображений коленного сустава, включая классификацию степени ОА "
        "по шкале Kellgren–Lawrence и прогнозирование прогрессии заболевания."
    )
    set_run(r2)

    # =====================================================================
    # ОБЗОР ЛИТЕРАТУРЫ
    # =====================================================================
    add_heading_text(doc, "Обзор литературы")

    add_paragraph(doc,
        "Метод Otsu [4] позволяет автоматически определять порог бинаризации, но имеет "
        "ограниченную точность на мягких тканях. U-Net [1] достигает Dice 0.85–0.90 на "
        "хрящевой ткани [13], nnU-Net [2] — 0.88–0.95, TransUNet [3] интегрирует Vision "
        "Transformer. Текстурные признаки GLCM [9] дополняют интенсивностный анализ."
    )

    add_paragraph(doc,
        "Для классификации степени остеоартрита используется шкала Kellgren–Lawrence [5] "
        "(KL 0–4). Методы Random Forest [8] и SVM применяются для классификации по "
        "KL-грейду [17]. В данной работе предлагается комплексный подход, объединяющий "
        "EDA, базовую сегментацию и модели ML."
    )

    # =====================================================================
    # МЕТОДЫ И МОДЕЛИ
    # =====================================================================
    add_heading_text(doc, "Методы и модели исследования")

    add_paragraph(doc,
        "Использован набор данных 3D Knee MRI Cartilage Segmentation (Kaggle/OAI): "
        "1 733 МРТ-среза с разметкой 5 классов (фон, бедренный хрящ, большеберцовый "
        "хрящ, костная ткань, дефекты хряща). Дисбаланс: KL-0–KL-3 по ~400 объектов "
        "(23%), KL-4 — 133 объекта (7.67%)."
    )

    add_table(doc,
        ["Класс", "Код", "Описание"],
        [
            ["Фон", "0", "Нетканевые области"],
            ["Бедренный хрящ", "1", "Хрящевая ткань бедренной кости"],
            ["Большеберцовый хрящ", "2", "Хрящевая ткань большеберцовой кости"],
            ["Костная ткань", "3", "Бедренная и большеберцовая кости"],
            ["Дефекты хряща", "4", "Области повреждения хрящевой ткани"],
        ],
        title="Таблица 1 — Анатомические классы сегментации"
    )

    add_paragraph(doc,
        "Из каждого МРТ-среза извлечены 20 интерпретируемых признаков [9, 12]: "
        "интенсивностные (mean, std, skew, kurtosis), спектральные (fft_mean), "
        "резкость (laplacian_var), границы (edge_density), текстурные GLCM "
        "(contrast, homogeneity, energy, correlation, entropy), морфологические "
        "(bone_area_ratio, cartilage_area_ratio, joint_space_width), клинические "
        "(osteophyte_score, sclerosis_index) и гистограммные (mean_gradient, "
        "hist_peak_pos, hist_spread)."
    )

    add_paragraph(doc,
        "Предварительная обработка включает удаление пропусков, нормализацию "
        "StandardScaler (x' = (x − μ) / σ) и стратифицированное разбиение 80/20. "
        "Нормализация необходима для линейных моделей, так как масштабы признаков "
        "существенно различаются: mean_intensity ∈ [90, 150], laplacian_var ∈ [600, 1200]."
    )

    add_paragraph(doc,
        "Базовая сегментация реализована конвейером [4]: нормализация → порог Otsu → "
        "морфологическое закрытие → анализ связных компонент → присвоение классов "
        "по размеру компонент. Данный подход обеспечивает Dice 0.20–0.35, что "
        "существенно ниже нейросетевых методов [1, 2], но служит baseline."
    )

    add_paragraph(doc,
        "Для классификации по KL-грейду обучены 5 моделей [7, 8]: Logistic Regression, "
        "Random Forest (300 деревьев), Decision Tree, SVM (RBF-ядро) и KNN (k=7). "
        "Для прогнозирования прогрессии ОА построена линейная регрессия на основе "
        "клинически обоснованных признаков [18, 19]."
    )

    # =====================================================================
    # РЕЗУЛЬТАТЫ
    # =====================================================================
    add_heading_text(doc, "Результаты вычислительного эксперимента")

    add_paragraph(doc,
        "Анализ распределения признаков показал статистически значимые различия "
        "между классами KL-грейдов. Плотность границ (edge_density) снижается с "
        "увеличением тяжести ОА: деградация хрящевой ткани приводит к потере "
        "чётких анатомических границ. Дисперсия лапласиана (laplacian_var) также "
        "уменьшается при прогрессировании заболевания."
    )

    add_paragraph(doc,
        "Корреляционный анализ выявил наибольшую связь с KL-грейдом у признаков "
        "joint_space_width, osteophyte_score и cartilage_area_ratio. Внутри группы "
        "текстурных признаков GLCM обнаружена мультиколлинеарность: корреляция между "
        "energy и homogeneity достигает r ≈ 0.9. PCA-проекция показала частичное "
        "перекрытие классов KL-0 — KL-4, что указывает на необходимость нелинейных моделей."
    )

    add_figure(doc,
        FIGURES_DIR / "class_distribution.png",
        "Рисунок 1 — Распределение классов по степеням остеоартрита (KL-grade)",
        width=Inches(3.2))

    add_table(doc,
        ["Модель", "Accuracy", "Precision", "Recall", "F1"],
        [
            ["Logistic Regression", "75.22%", "75.61%", "75.22%", "75.19%"],
            ["SVM", "74.06%", "74.64%", "74.06%", "73.96%"],
            ["Random Forest", "71.47%", "73.69%", "71.47%", "70.36%"],
            ["KNN", "63.11%", "63.33%", "63.11%", "61.60%"],
            ["Decision Tree", "54.18%", "54.34%", "54.18%", "54.19%"],
        ],
        title="Таблица 2 — Результаты классификации по KL-грейду"
    )

    add_paragraph(doc,
        "Logistic Regression показала лучший результат (F1 = 75.19%). SVM — "
        "сопоставимый (74.06%). Decision Tree переобучился на 20 признаках (54.18%)."
    )

    add_figure(doc,
        FIGURES_DIR / "confusion_matrices.png",
        "Рисунок 2 — Confusion matrices: Logistic Regression vs Random Forest",
        width=Inches(3.2))

    add_paragraph(doc,
        "Анализ важности признаков модели Random Forest показал, что наибольший "
        "вклад в классификацию вносят клинически значимые морфологические признаки: "
        "joint_space_width (ширина суставной щели), osteophyte_score (остеофиты) "
        "и cartilage_area_ratio (доля хрящевой ткани). Базовые статистики "
        "интенсивности оказались наименее значимыми."
    )

    add_figure(doc,
        FIGURES_DIR / "feature_importance.png",
        "Рисунок 3 — Важность признаков модели Random Forest",
        width=Inches(3.2))

    add_table(doc,
        ["Метрика", "Значение", "Интерпретация"],
        [
            ["R²", "0.8492", "Модель объясняет 84.9% вариации"],
            ["RMSE", "±0.37%", "Средняя ошибка предсказания"],
            ["MAE", "0.29%", "Средняя абсолютная ошибка"],
        ],
        title="Таблица 3 — Метрики модели линейной регрессии"
    )

    add_paragraph(doc,
        "Ключевые факторы прогрессии: osteophyte_score (+0.38), sclerosis_index "
        "(+0.34), joint_space_width (−0.32). Сужение суставной щели и выраженность "
        "остеофитов — главные предикторы ускоренной прогрессии ОА."
    )

    # =====================================================================
    # ОБСУЖДЕНИЕ
    # =====================================================================
    add_heading_text(doc, "Обсуждение")

    add_paragraph(doc,
        "Logistic Regression достигает наибольшей точности (75.22%), близкий "
        "результат SVM (74.06%) подтверждает линейную разделимость после нормализации. "
        "Ключевые предикторы — ширина суставной щели, выраженность остеофитов и доля "
        "хрящевой ткани — согласуются с клиническими наблюдениями [5, 18]. "
        "Ограничения: дисбаланс класса KL-4 (7.67%), моделированная целевая переменная "
        "регрессии, низкий Dice базовой сегментации (0.20–0.35). Перспективы: реализация "
        "U-Net/nnU-Net [1, 2], 3D-сегментация и валидация на данных OAI [6]."
    )

    # =====================================================================
    # ЗАКЛЮЧЕНИЕ
    # =====================================================================
    add_heading_text(doc, "Заключение")

    add_paragraph(doc,
        "В данной работе рассмотрена задача автоматической сегментации и анализа "
        "МРТ-изображений коленного сустава. Проведён разведочный анализ с извлечением "
        "20 признаков, реализована базовая сегментация (Dice 0.20–0.35) и построены "
        "модели классификации (Logistic Regression: Accuracy = 75.22%, F1 = 75.19%) "
        "и регрессии (R² = 0.85, RMSE = ±0.37%). В дальнейшем планируется реализация "
        "U-Net/nnU-Net, переход к 3D-анализу и валидация на данных OAI."
    )

    # =====================================================================
    # СПИСОК ЛИТЕРАТУРЫ
    # =====================================================================
    add_heading_text(doc, "Список литературы")

    references = [
        "Ronneberger O., Fischer P., Brox T. U-Net: Convolutional Networks for "
        "Biomedical Image Segmentation // MICCAI, 2015.",

        "Isensee F., Jaeger P., Kohl S. et al. nnU-Net: a self-configuring method "
        "for deep learning-based biomedical image segmentation // Nature Methods, 2021.",

        "Chen J., Lu Y., Yu Q. et al. TransUNet: Transformers Make Strong Encoders "
        "for Medical Image Segmentation // arXiv:2102.04306, 2021.",

        "Otsu N. A threshold selection method from gray-level histograms // IEEE Trans. "
        "Systems, Man, and Cybernetics, 1979.",

        "Kellgren J.H., Lawrence J.S. Radiological assessment of osteo-arthrosis // "
        "Annals of the Rheumatic Diseases, 1957.",

        "Osteoarthritis Initiative (OAI). NIAMS. https://nda.nih.gov/oai/",

        "Pedregosa F. et al. Scikit-learn: Machine Learning in Python // JMLR, 2011.",

        "Breiman L. Random Forests // Machine Learning, 2001.",

        "Haralick R.M. et al. Textural Features for Image Classification // IEEE Trans. "
        "SMC, 1973.",

        "McKinney W. Data Structures for Statistical Computing in Python // "
        "Proc. 9th Python in Science Conf., 2010.",

        "Hunter J.D. Matplotlib: A 2D Graphics Environment // Computing in Science "
        "& Engineering, 2007.",

        "Bradski G. The OpenCV Library // Dr. Dobb's Journal, 2000.",

        "Norman B., Pedoia V., Majumdar S. Use of 2D U-Net for automated cartilage "
        "and meniscus segmentation of knee MR imaging data // Radiology, 2018.",

        "Ambellan F. et al. Automated segmentation of knee bone and cartilage combining "
        "statistical shape knowledge and CNNs // Medical Image Analysis, 2019.",

        "Liu F. et al. Deep CNN and 3D deformable approach for tissue segmentation in "
        "musculoskeletal MRI // Magnetic Resonance in Medicine, 2018.",

        "Panfilov E. et al. Deep learning-based segmentation of knee MRI for automatic "
        "morphological assessment of cartilage // Osteoarthritis and Cartilage, 2022.",

        "Tiulpin A. et al. Automatic knee osteoarthritis diagnosis from plain "
        "radiographs: A deep learning-based approach // Scientific Reports, 2018.",

        "Eckstein F., Wirth W., Nevitt M. Recent advances in osteoarthritis imaging — "
        "the OAI // Nature Reviews Rheumatology, 2012.",

        "Wirth W., Eckstein F. A technique for regional analysis of femorotibial "
        "cartilage thickness based on quantitative MRI // IEEE Trans. Med. Imaging, 2008.",

        "3D Knee MRI Cartilage Segmentation Dataset // Kaggle.",
    ]

    for i, ref in enumerate(references):
        p = add_paragraph(doc, f"{i+1}. {ref}", font_size=14,
                          first_line_indent=Cm(1.0))
        p.paragraph_format.space_after = Pt(2)

    # =====================================================================
    # SAVE
    # =====================================================================
    output_path = BASE_DIR / "Маратов_Ерасыл_Балканович_Статья.docx"
    doc.save(str(output_path))
    print(f"Article saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    build_article()
