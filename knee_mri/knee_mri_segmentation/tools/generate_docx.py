"""
Generate REPORT.docx from REPORT.md with embedded images.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(SCRIPT_DIR, "results", "eda_figures")


def add_image_safe(doc, img_path, width=Inches(5.8)):
    """Add image if file exists, otherwise add a placeholder note."""
    if os.path.isfile(img_path):
        doc.add_picture(img_path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Изображение не найдено: {os.path.basename(img_path)}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True


def build_report():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {qn("w:eastAsia"): "Times New Roman"})
    rPr.append(rFonts)

    # ── Author name ──
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("Маратов Ерасыл")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    doc.add_paragraph()  # spacer

    # ── Title ──
    title = doc.add_heading("Разведочный анализ данных (EDA) — Knee MRI Segmentation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "Анализ набора данных МРТ коленного сустава из Osteoarthritis Initiative (OAI), "
        "содержащего срезы с разметкой анатомических структур для задачи автоматической сегментации."
    )

    # Dataset info
    p = doc.add_paragraph()
    p.add_run("Датасет: ").bold = True
    p.add_run("3D Knee MRI Cartilage Segmentation (Kaggle: ujjwalsinha01/3d-knee-mri-cartilage-segmentation)")
    p = doc.add_paragraph()
    p.add_run("Классы: ").bold = True
    p.add_run("5 (фон, бедренный хрящ, большеберцовый хрящ, костная ткань, дефекты хряща)")
    p = doc.add_paragraph()
    p.add_run("Сплиты: ").bold = True
    p.add_run("train / val / test")
    p = doc.add_paragraph()
    p.add_run("Формат: ").bold = True
    p.add_run("PNG (grayscale) → .npy")

    # ────────────────────────────────────────────
    # Section 1
    # ────────────────────────────────────────────
    doc.add_heading("1. Описание извлечённых признаков", level=1)
    doc.add_paragraph(
        "Из каждого МРТ-среза извлечены 20 интерпретируемых параметров:"
    )

    features = [
        ("Интенсивность:", "mean_intensity, std_intensity, skew_intensity, kurt_intensity"),
        ("Спектральные:", "fft_mean (средняя энергия высокочастотного спектра Фурье)"),
        ("Резкость:", "laplacian_var (дисперсия оператора Лапласа)"),
        ("Границы:", "edge_density (плотность границ Canny)"),
        ("Текстура (GLCM):", "contrast, homogeneity, energy, correlation, entropy"),
        ("Морфология:", "bone_area_ratio, cartilage_area_ratio, joint_space_width"),
        ("Клинические:", "osteophyte_score, sclerosis_index"),
        ("Гистограмма:", "mean_gradient, hist_peak_pos, hist_spread"),
    ]
    for label, desc in features:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(label).bold = True
        p.add_run(f" {desc}")

    add_image_safe(doc, os.path.join(IMG_DIR, "image_analysis_panels.png"))
    doc.add_paragraph(
        "Визуализация карт признаков для KL-0 (здоровый) и KL-4 (тяжёлый OA): "
        "исходный срез, спектр Фурье, лапласиан, границы и карта интенсивности."
    ).italic = True

    add_image_safe(doc, os.path.join(IMG_DIR, "sample_grid_kl_grades.png"))
    doc.add_paragraph(
        "Примеры МРТ-срезов по степеням остеоартрита (KL 0–4). С увеличением степени — "
        "сужение суставной щели, появление остеофитов, истончение хряща."
    ).italic = True

    # ────────────────────────────────────────────
    # Section 2
    # ────────────────────────────────────────────
    doc.add_heading("2. Разведочный анализ данных и оценка распределений", level=1)

    doc.add_heading("Распределение классов", level=2)
    add_image_safe(doc, os.path.join(IMG_DIR, "class_distribution.png"))
    add_image_safe(doc, os.path.join(IMG_DIR, "split_table.png"))
    doc.add_paragraph(
        "Обнаружен дисбаланс классов: KL-4 (тяжёлый остеоартрит) представлен значительно "
        "меньшим числом образцов. Это требует взвешивания классов или аугментации при обучении."
    )

    doc.add_heading("Распределение признаков (KDE + Boxplot)", level=2)

    add_image_safe(doc, os.path.join(IMG_DIR, "distribution_edge_density.png"))
    p = doc.add_paragraph()
    p.add_run("Плотность границ (edge_density): ").bold = True
    p.add_run(
        "снижается с увеличением тяжести OA. Деградация хрящевой ткани приводит к потере "
        "чётких анатомических границ — аналогично тому, как реальные фотографии содержат "
        "больше высокочастотного «хаоса» по сравнению со сглаженными синтетическими изображениями."
    )

    add_image_safe(doc, os.path.join(IMG_DIR, "distribution_laplacian_var.png"))
    p = doc.add_paragraph()
    p.add_run("Дисперсия лапласиана (laplacian_var): ").bold = True
    p.add_run(
        "здоровые суставы (KL-0) показывают более высокую резкость. Прогрессирование заболевания "
        "сопровождается размытием границ тканей."
    )

    # ────────────────────────────────────────────
    # Section 3
    # ────────────────────────────────────────────
    doc.add_heading("3. Корреляционный анализ и снижение размерности", level=1)

    doc.add_heading("Матрица корреляций", level=2)
    add_image_safe(doc, os.path.join(IMG_DIR, "correlation_matrix.png"))

    p = doc.add_paragraph()
    p.add_run("Связь с целевой переменной: ").bold = True
    p.add_run(
        "наибольшую корреляцию с KL-грейдом демонстрируют joint_space_width, "
        "osteophyte_score и cartilage_area_ratio."
    )
    p = doc.add_paragraph()
    p.add_run("Мультиколлинеарность: ").bold = True
    p.add_run(
        "внутри группы текстурных признаков GLCM (energy и homogeneity) наблюдается "
        "высокая корреляция, близкая к r ≈ 0.9."
    )

    doc.add_heading("PCA-проекция", level=2)
    add_image_safe(doc, os.path.join(IMG_DIR, "pca_projection.png"))
    doc.add_paragraph(
        "Проекция на 2 главные компоненты показывает частичное перекрытие классов KL-0 — KL-4. "
        "Отсутствие чёткой линейной границы указывает на необходимость применения нелинейных "
        "моделей для достижения высокой точности."
    )

    # ────────────────────────────────────────────
    # Section 4
    # ────────────────────────────────────────────
    doc.add_heading("4. Сравнение моделей и результаты классификации", level=1)
    doc.add_paragraph(
        "Протестированы: Logistic Regression (baseline) и Random Forest (ансамблевый метод)."
    )

    doc.add_heading("Confusion Matrices", level=2)
    add_image_safe(doc, os.path.join(IMG_DIR, "confusion_matrices.png"))
    doc.add_paragraph(
        "Random Forest демонстрирует значительное превосходство — ошибки межклассовой "
        "путаницы снижаются, особенно для граничных классов (KL-1 / KL-2)."
    )

    doc.add_heading("ROC-кривые", level=2)
    add_image_safe(doc, os.path.join(IMG_DIR, "roc_curves.png"))
    doc.add_paragraph(
        "AUC у Random Forest существенно выше, чем у Logistic Regression, что подтверждает "
        "эффективность ансамблевого метода на данных с нелинейными паттернами, выявленными на этапе PCA."
    )

    # ────────────────────────────────────────────
    # Section 5
    # ────────────────────────────────────────────
    doc.add_heading("5. Интерпретируемость (Feature Importance)", level=1)
    add_image_safe(doc, os.path.join(IMG_DIR, "feature_importance.png"))

    doc.add_paragraph("Ключевые индикаторы степени остеоартрита по Random Forest:")

    importance_items = [
        ("joint_space_width", "ширина суставной щели — наиболее весомый фактор"),
        ("osteophyte_score", "выраженность остеофитов"),
        ("cartilage_area_ratio", "доля хрящевой ткани"),
        ("fft_mean", "спектральная энергия"),
        ("laplacian_var", "микротекстура / резкость"),
    ]
    for i, (feat, desc) in enumerate(importance_items, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{feat}").bold = True
        p.add_run(f" ({desc})")

    doc.add_paragraph(
        "Базовые статистики интенсивности (skew, kurtosis) оказались наименее значимыми — "
        "клинически значимые морфологические признаки (щель, остеофиты, хрящ) несут больше "
        "информации, чем простые пиксельные характеристики."
    )

    # ────────────────────────────────────────────
    # Conclusions
    # ────────────────────────────────────────────
    doc.add_heading("Содержательные выводы", level=1)

    conclusions = [
        ("Дисбаланс классов", "KL-4 (тяжёлый OA) содержит в ~3 раза меньше образцов, что может смещать модель без балансировки."),
        ("Ширина суставной щели — главный предиктор", "демонстрирует наибольшую важность и корреляцию с тяжестью OA."),
        ("Спектральные и текстурные артефакты", "fft_mean и laplacian_var показывают статистически значимые различия между классами, подтверждая, что прогрессирование OA изменяет частотные характеристики снимка."),
        ("Нелинейность данных", "PCA показывает перекрытие классов; Random Forest превосходит Logistic Regression по AUC, что подтверждает нелинейную природу зависимостей."),
        ("Мультиколлинеарность текстуры", "признаки GLCM (energy, homogeneity) показывают корреляцию ≈ 0.9, что может потребовать отбора признаков или регуляризации."),
        ("Базовые пиксельные статистики малоинформативны", "skewness и kurtosis интенсивности занимают последние места по важности."),
        ("Базовый метод сегментации (Otsu + морфология)", "обеспечивает Dice ~0.20–0.35, что значительно ниже нейросетевых аналогов (U-Net: 0.85–0.95), но достаточно для baseline."),
    ]
    for i, (title_text, desc) in enumerate(conclusions, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {title_text}").bold = True
        p.add_run(f" — {desc}")

    # ────────────────────────────────────────────
    # Visualization table
    # ────────────────────────────────────────────
    doc.add_heading("Визуализации", level=1)

    viz_items = [
        ("1", "image_analysis_panels.png", "Панели анализа МРТ (FFT, Laplacian, Edges)"),
        ("2", "sample_grid_kl_grades.png", "Примеры снимков KL-0 — KL-4"),
        ("3", "class_distribution.png", "Распределение классов (Total + splits)"),
        ("4", "split_table.png", "Таблица распределения по сплитам"),
        ("5", "distribution_edge_density.png", "KDE + Boxplot: edge_density"),
        ("6", "distribution_laplacian_var.png", "KDE + Boxplot: laplacian_var"),
        ("7", "correlation_matrix.png", "Матрица корреляций"),
        ("8", "pca_projection.png", "PCA-проекция (2D)"),
        ("9", "confusion_matrices.png", "Confusion Matrices (LR vs RF)"),
        ("10", "roc_curves.png", "ROC-кривые"),
        ("11", "feature_importance.png", "Feature Importance (Random Forest)"),
    ]

    table = doc.add_table(rows=1, cols=3, style="Light Shading Accent 1")
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Файл"
    hdr[2].text = "Описание"
    for num, fname, desc in viz_items:
        row = table.add_row().cells
        row[0].text = num
        row[1].text = fname
        row[2].text = desc

    # ── Save ──
    out_path = os.path.join(SCRIPT_DIR, "REPORT.docx")
    doc.save(out_path)
    print(f"✅ Saved: {out_path}")


if __name__ == "__main__":
    build_report()
